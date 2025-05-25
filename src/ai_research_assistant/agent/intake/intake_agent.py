"""
Intake Agent for AI Research Assistant

This agent specializes in document processing, case classification, and initial analysis.
It serves as the entry point for new documents and cases into the research system.
"""

import logging
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# Document processing imports
try:
    from PyPDF2 import PdfReader

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx  # type: ignore

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# Core agent imports
from ..core.base_agent import AgentConfig, AgentTask, BaseAgent

logger = logging.getLogger(__name__)


class DocumentType:
    """Document type classifications"""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    JSON = "json"
    UNKNOWN = "unknown"


class CaseType:
    """Case type classifications for legal/research documents"""

    LEGAL_CASE = "legal_case"
    RESEARCH_PAPER = "research_paper"
    TECHNICAL_DOC = "technical_doc"
    BUSINESS_DOC = "business_doc"
    PERSONAL_DOC = "personal_doc"
    UNKNOWN = "unknown"


class IntakeAgent(BaseAgent):
    """
    Specialized agent for document intake, processing, and classification.

    Capabilities:
    - Document format detection and validation
    - Text extraction from multiple file formats
    - Metadata extraction and enrichment
    - Content classification and categorization
    - Initial similarity analysis
    - Database indexing preparation
    """

    def __init__(
        self,
        agent_id: Optional[str] = None,
        config: Optional[AgentConfig] = None,
        global_settings_manager=None,
        intake_directory: str = "./tmp/intake",
        processed_directory: str = "./tmp/processed",
        max_file_size_mb: int = 100,
        **kwargs,
    ):
        super().__init__(
            agent_id=agent_id,
            config=config,
            global_settings_manager=global_settings_manager,
            **kwargs,
        )

        # Intake-specific configuration
        self.intake_directory = Path(intake_directory)
        self.processed_directory = Path(processed_directory)
        self.max_file_size_mb = max_file_size_mb

        # Create directories if they don't exist
        self.intake_directory.mkdir(parents=True, exist_ok=True)
        self.processed_directory.mkdir(parents=True, exist_ok=True)

        # Document processing stats
        self.documents_processed = 0
        self.classification_cache: Dict[str, str] = {}

        logger.info(f"IntakeAgent initialized: {self.agent_id}")

    def get_supported_task_types(self) -> List[str]:
        """Return list of task types this agent can handle."""
        return [
            "process_document",
            "classify_document",
            "extract_metadata",
            "validate_file",
            "batch_process",
            "similarity_check",
        ]

    async def execute_task(self, task: AgentTask) -> Dict[str, Any]:
        """Execute a specific intake task."""
        task_type = task.task_type
        params = task.parameters

        try:
            if task_type == "process_document":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._process_document(
                    file_path=file_path,
                    document_id=params.get("document_id"),
                    metadata=params.get("metadata", {}),
                )

            elif task_type == "classify_document":
                content = params.get("content")
                if not content:
                    raise ValueError("content parameter is required")
                return await self._classify_document(
                    content=content, file_path=params.get("file_path")
                )

            elif task_type == "extract_metadata":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._extract_metadata(file_path=Path(file_path))

            elif task_type == "validate_file":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._validate_file(file_path=file_path)

            elif task_type == "batch_process":
                return await self._batch_process(
                    directory=params.get("directory", self.intake_directory),
                    file_patterns=params.get(
                        "file_patterns", ["*.pdf", "*.docx", "*.txt"]
                    ),
                )

            elif task_type == "similarity_check":
                content = params.get("content")
                document_id = params.get("document_id")
                if not content or not document_id:
                    raise ValueError("content and document_id parameters are required")
                return await self._similarity_check(
                    content=content, document_id=document_id
                )

            else:
                raise ValueError(f"Unsupported task type: {task_type}")

        except Exception as e:
            logger.error(f"Task execution failed for {task_type}: {e}", exc_info=True)
            raise

    async def _process_document(
        self,
        file_path: Union[str, Path],
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Process a single document through the full intake pipeline."""
        file_path = Path(file_path)
        document_id = document_id or str(uuid.uuid4())

        logger.info(f"Processing document: {file_path}")

        # Step 1: Validate file
        validation_result = await self._validate_file(file_path)
        if not validation_result["valid"]:
            return {
                "success": False,
                "document_id": document_id,
                "error": validation_result["error"],
                "stage": "validation",
            }

        # Step 2: Extract text content
        content_result = await self._extract_text_content(file_path)
        if not content_result["success"]:
            return {
                "success": False,
                "document_id": document_id,
                "error": content_result["error"],
                "stage": "content_extraction",
            }

        # Step 3: Extract metadata
        metadata_result = await self._extract_metadata(file_path)

        # Step 4: Classify document
        classification_result = await self._classify_document(
            content=content_result["content"], file_path=file_path
        )

        # Step 5: Perform similarity check if requested
        similarity_result = None
        if len(content_result["content"]) > 100:  # Only for substantial content
            similarity_result = await self._similarity_check(
                content=content_result["content"], document_id=document_id
            )

        # Compile final result
        result = {
            "success": True,
            "document_id": document_id,
            "file_path": str(file_path),
            "document_type": validation_result["document_type"],
            "file_size_mb": validation_result["file_size_mb"],
            "content": content_result["content"],
            "content_length": len(content_result["content"]),
            "metadata": {
                **metadata_result.get("metadata", {}),
                **(metadata or {}),
                "processed_at": datetime.now().isoformat(),
                "processed_by": self.agent_id,
            },
            "classification": classification_result,
            "similarity": similarity_result,
        }

        # Move file to processed directory
        processed_path = self.processed_directory / f"{document_id}_{file_path.name}"
        try:
            if file_path != processed_path:
                import shutil

                shutil.copy2(file_path, processed_path)
                result["processed_path"] = str(processed_path)
        except Exception as e:
            logger.warning(f"Failed to move file to processed directory: {e}")

        self.documents_processed += 1
        logger.info(f"Document processed successfully: {document_id}")

        return result

    async def _validate_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Validate file format, size, and accessibility."""
        file_path = Path(file_path)

        try:
            if not file_path.exists():
                return {
                    "valid": False,
                    "error": "File does not exist",
                    "file_path": str(file_path),
                }

            if not file_path.is_file():
                return {
                    "valid": False,
                    "error": "Path is not a file",
                    "file_path": str(file_path),
                }

            # Check file size
            file_size_bytes = file_path.stat().st_size
            file_size_mb = file_size_bytes / (1024 * 1024)

            if file_size_mb > self.max_file_size_mb:
                return {
                    "valid": False,
                    "error": f"File too large: {file_size_mb:.2f}MB > {self.max_file_size_mb}MB",
                    "file_size_mb": file_size_mb,
                }

            # Determine document type
            document_type = self._determine_document_type(file_path)

            return {
                "valid": True,
                "file_path": str(file_path),
                "file_size_mb": file_size_mb,
                "document_type": document_type,
            }

        except Exception as e:
            return {
                "valid": False,
                "error": f"File validation error: {str(e)}",
                "file_path": str(file_path),
            }

    def _determine_document_type(self, file_path: Path) -> str:
        """Determine document type based on file extension and content."""
        extension = file_path.suffix.lower()

        type_mapping = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".json": DocumentType.JSON,
        }

        return type_mapping.get(extension, DocumentType.UNKNOWN)

    async def _extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text content from various file formats."""
        document_type = self._determine_document_type(file_path)

        try:
            if document_type == DocumentType.PDF:
                return await self._extract_pdf_content(file_path)
            elif document_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            elif document_type == DocumentType.TXT:
                return await self._extract_txt_content(file_path)
            elif document_type == DocumentType.HTML:
                return await self._extract_html_content(file_path)
            elif document_type == DocumentType.JSON:
                return await self._extract_json_content(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document type: {document_type}",
                    "content": "",
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"Content extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files."""
        if not HAS_PYPDF2:
            return {
                "success": False,
                "error": "PyPDF2 not available. Install with: pip install PyPDF2",
                "content": "",
            }

        try:
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                content = ""

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )

                return {
                    "success": True,
                    "content": content.strip(),
                    "page_count": len(reader.pages),
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        if not HAS_PYTHON_DOCX:
            return {
                "success": False,
                "error": "python-docx not available. Install with: pip install python-docx",
                "content": "",
            }

        try:
            doc = docx.Document(file_path)
            content = ""

            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            return {
                "success": True,
                "content": content.strip(),
                "paragraph_count": len(doc.paragraphs),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_txt_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()

            return {
                "success": True,
                "content": content,
                "character_count": len(content),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_html_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from HTML files."""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                soup = BeautifulSoup(file.read(), "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text()

            return {
                "success": True,
                "content": content.strip(),
                "html_title": soup.title.string if soup.title else None,
            }
        except ImportError:
            return {
                "success": False,
                "error": "BeautifulSoup not available. Install with: pip install beautifulsoup4",
                "content": "",
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"HTML extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_json_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from JSON files."""
        try:
            import json

            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Convert JSON to readable text
            content = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                "success": True,
                "content": content,
                "json_keys": list(data.keys()) if isinstance(data, dict) else None,
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"JSON extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file."""
        try:
            stat = file_path.stat()

            metadata = {
                "file_name": file_path.name,
                "file_extension": file_path.suffix,
                "file_size_bytes": stat.st_size,
                "file_size_mb": stat.st_size / (1024 * 1024),
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_time": datetime.fromtimestamp(stat.st_atime).isoformat(),
            }

            return {"success": True, "metadata": metadata}

        except Exception as e:
            return {
                "success": False,
                "error": f"Metadata extraction failed: {str(e)}",
                "metadata": {},
            }

    async def _classify_document(
        self, content: str, file_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Classify document type and content category using LLM analysis."""
        if not content or len(content.strip()) < 50:
            return {
                "case_type": CaseType.UNKNOWN,
                "confidence": 0.0,
                "reasoning": "Insufficient content for classification",
            }

        try:
            # Use LLM for intelligent classification
            classification_prompt = f"""
            Analyze the following document content and classify it into one of these categories:
            1. legal_case - Legal documents, court cases, legal briefs
            2. research_paper - Academic papers, research studies, scientific documents
            3. technical_doc - Technical documentation, manuals, specifications
            4. business_doc - Business reports, proposals, correspondence
            5. personal_doc - Personal documents, letters, notes
            6. unknown - Cannot be classified

            Content to analyze:
            {content[:2000]}...

            Respond with JSON format:
            {{
                "case_type": "category_name",
                "confidence": 0.0-1.0,
                "reasoning": "Brief explanation",
                "key_indicators": ["list", "of", "indicators"]
            }}
            """

            response = await self.llm.ainvoke(classification_prompt)

            # Parse LLM response
            try:
                import json

                result = json.loads(response.content)

                # Validate classification
                valid_types = [
                    CaseType.LEGAL_CASE,
                    CaseType.RESEARCH_PAPER,
                    CaseType.TECHNICAL_DOC,
                    CaseType.BUSINESS_DOC,
                    CaseType.PERSONAL_DOC,
                    CaseType.UNKNOWN,
                ]

                if result.get("case_type") not in valid_types:
                    result["case_type"] = CaseType.UNKNOWN
                    result["confidence"] = 0.0

                return result

            except (json.JSONDecodeError, KeyError):
                # Fallback to simple keyword-based classification
                return self._fallback_classify(content, file_path)

        except Exception as e:
            logger.warning(f"LLM classification failed: {e}")
            return self._fallback_classify(content, file_path)

    def _fallback_classify(
        self, content: str, file_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Fallback classification using simple keyword matching."""
        content_lower = content.lower()

        # Legal indicators
        legal_keywords = [
            "court",
            "plaintiff",
            "defendant",
            "judgment",
            "appeal",
            "statute",
            "legal",
            "law",
            "case",
            "trial",
            "hearing",
            "docket",
        ]
        legal_score = sum(1 for keyword in legal_keywords if keyword in content_lower)

        # Research indicators
        research_keywords = [
            "abstract",
            "methodology",
            "conclusion",
            "references",
            "study",
            "research",
            "analysis",
            "experiment",
            "hypothesis",
            "data",
        ]
        research_score = sum(
            1 for keyword in research_keywords if keyword in content_lower
        )

        # Technical indicators
        tech_keywords = [
            "specification",
            "manual",
            "documentation",
            "technical",
            "system",
            "configuration",
            "implementation",
            "architecture",
            "api",
            "protocol",
        ]
        tech_score = sum(1 for keyword in tech_keywords if keyword in content_lower)

        # Business indicators
        business_keywords = [
            "proposal",
            "report",
            "meeting",
            "project",
            "budget",
            "revenue",
            "business",
            "management",
            "strategy",
            "market",
            "client",
        ]
        business_score = sum(
            1 for keyword in business_keywords if keyword in content_lower
        )

        # Determine classification based on highest score
        scores = {
            CaseType.LEGAL_CASE: legal_score,
            CaseType.RESEARCH_PAPER: research_score,
            CaseType.TECHNICAL_DOC: tech_score,
            CaseType.BUSINESS_DOC: business_score,
        }

        max_score = max(scores.values())
        if max_score == 0:
            return {
                "case_type": CaseType.UNKNOWN,
                "confidence": 0.0,
                "reasoning": "No clear indicators found",
                "key_indicators": [],
            }

        case_type = max(scores, key=lambda k: scores[k])
        confidence = min(max_score / 10.0, 1.0)  # Normalize to 0-1

        return {
            "case_type": case_type,
            "confidence": confidence,
            "reasoning": f"Keyword-based classification (score: {max_score})",
            "key_indicators": [k for k, v in scores.items() if v > 0],
        }

    async def _similarity_check(self, content: str, document_id: str) -> Dict[str, Any]:
        """Perform basic similarity analysis with existing documents."""
        # This is a placeholder for similarity analysis
        # In a full implementation, this would compare against a database
        # of existing documents using embedding similarity

        try:
            # Simple similarity check based on content length and structure
            word_count = len(content.split())
            unique_words = len(set(content.lower().split()))

            similarity_metrics = {
                "word_count": word_count,
                "unique_words": unique_words,
                "vocabulary_richness": unique_words / max(word_count, 1),
                "similar_documents": [],  # Would be populated from database search
                "similarity_score": 0.0,  # Would be calculated against existing docs
            }

            return {
                "success": True,
                "document_id": document_id,
                "metrics": similarity_metrics,
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Similarity check failed: {str(e)}",
                "document_id": document_id,
            }

    async def _batch_process(
        self, directory: Union[str, Path], file_patterns: List[str]
    ) -> Dict[str, Any]:
        """Process multiple documents in a directory."""
        directory = Path(directory)

        if not directory.exists():
            return {
                "success": False,
                "error": f"Directory does not exist: {directory}",
                "processed_count": 0,
            }

        results = []
        processed_count = 0
        error_count = 0

        try:
            for pattern in file_patterns:
                for file_path in directory.glob(pattern):
                    try:
                        result = await self._process_document(file_path)
                        results.append(result)

                        if result["success"]:
                            processed_count += 1
                        else:
                            error_count += 1

                    except Exception as e:
                        error_count += 1
                        results.append(
                            {
                                "success": False,
                                "file_path": str(file_path),
                                "error": str(e),
                            }
                        )

            return {
                "success": True,
                "directory": str(directory),
                "processed_count": processed_count,
                "error_count": error_count,
                "total_files": len(results),
                "results": results,
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Batch processing failed: {str(e)}",
                "processed_count": processed_count,
                "error_count": error_count,
            }

    async def get_intake_statistics(self) -> Dict[str, Any]:
        """Get statistics about document intake operations."""
        return {
            "agent_id": self.agent_id,
            "documents_processed": self.documents_processed,
            "intake_directory": str(self.intake_directory),
            "processed_directory": str(self.processed_directory),
            "max_file_size_mb": self.max_file_size_mb,
            "supported_formats": [
                DocumentType.PDF,
                DocumentType.DOCX,
                DocumentType.TXT,
                DocumentType.HTML,
                DocumentType.JSON,
            ],
            "cache_size": len(self.classification_cache),
        }
