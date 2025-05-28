"""
Intake Agent for AI Research Assistant

This agent specializes in document processing, case classification, and initial analysis.
It serves as the entry point for new documents and cases into the research system.

Enhanced with:
- MCP Rust filesystem integration for high-speed file operations
- OCR capabilities with pytesseract
- Optimized processing order (OCR documents processed last)
- Proper LLM model selection (respects Google Gemini Pro settings)
"""

import logging
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# Document processing imports
try:
    from PyPDF2 import PdfReader

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx  # type: ignore

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# OCR imports
try:
    import cv2  # type: ignore
    import numpy as np
    import pytesseract  # type: ignore
    from PIL import Image

    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Core agent imports
from ..core.base_agent import AgentConfig, AgentTask, BaseAgent

logger = logging.getLogger(__name__)


class DocumentType:
    """Document type classifications"""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    JSON = "json"
    IMAGE = "image"  # Added for OCR support
    SCANNED_PDF = "scanned_pdf"  # Added for OCR support
    UNKNOWN = "unknown"


class CaseType:
    """Case type classifications for legal/research documents"""

    LEGAL_CASE = "legal_case"
    RESEARCH_PAPER = "research_paper"
    TECHNICAL_DOC = "technical_doc"
    BUSINESS_DOC = "business_doc"
    PERSONAL_DOC = "personal_doc"
    UNKNOWN = "unknown"


class DocumentSource:
    """Document source classifications for WCB tracking"""

    INSURANCE_COMPANY = "insurance_company"
    USER_CLAIMANT = "user_claimant"
    WCB_BOARD = "wcb_board"
    MEDICAL_PROVIDER = "medical_provider"
    EMPLOYER = "employer"
    LEGAL_REPRESENTATIVE = "legal_representative"
    UNKNOWN = "unknown"


class SubmissionType:
    """Types of submissions in WCB cases"""

    INITIAL_CLAIM = "initial_claim"
    MEDICAL_REPORT = "medical_report"
    DENIAL_RESPONSE = "denial_response"
    APPEAL_SUBMISSION = "appeal_submission"
    SUPPORTING_EVIDENCE = "supporting_evidence"
    COUNTER_ARGUMENT = "counter_argument"
    EXPERT_OPINION = "expert_opinion"
    CORRESPONDENCE = "correspondence"
    UNKNOWN = "unknown"


class WCBTracker:
    """Tracks WCB case submissions, doctors, and key points"""

    def __init__(self, case_id: str):
        self.case_id = case_id
        self.submissions = []
        self.insurance_doctors = set()
        self.user_doctors = set()
        self.medical_providers = set()
        self.key_points_made = []
        self.submission_timeline = []
        self.document_sources = {}

    def add_submission(self, document_info: Dict[str, Any]):
        """Add a new document submission to tracking"""
        self.submissions.append(document_info)
        self.submission_timeline.append(
            {
                "timestamp": document_info.get(
                    "submission_date", datetime.now().isoformat()
                ),
                "document": document_info.get("file_name"),
                "source": document_info.get("source"),
                "type": document_info.get("submission_type"),
            }
        )

    def add_medical_provider(self, doctor_name: str, source: str):
        """Add medical provider to appropriate tracking set"""
        if source == DocumentSource.INSURANCE_COMPANY:
            self.insurance_doctors.add(doctor_name)
        elif source == DocumentSource.USER_CLAIMANT:
            self.user_doctors.add(doctor_name)

        self.medical_providers.add(doctor_name)

    def add_key_point(self, point: str, document: str, source: str):
        """Add key point made in a document"""
        self.key_points_made.append(
            {
                "point": point,
                "document": document,
                "source": source,
                "timestamp": datetime.now().isoformat(),
            }
        )

    def get_summary(self) -> Dict[str, Any]:
        """Get comprehensive case tracking summary"""
        return {
            "case_id": self.case_id,
            "total_submissions": len(self.submissions),
            "insurance_doctors": list(self.insurance_doctors),
            "user_doctors": list(self.user_doctors),
            "total_medical_providers": len(self.medical_providers),
            "key_points_count": len(self.key_points_made),
            "submission_timeline": sorted(
                self.submission_timeline, key=lambda x: x["timestamp"]
            ),
            "sources_breakdown": self._get_sources_breakdown(),
        }

    def _get_sources_breakdown(self) -> Dict[str, int]:
        """Get breakdown of documents by source"""
        breakdown = {}
        for submission in self.submissions:
            source = submission.get("source", "unknown")
            breakdown[source] = breakdown.get(source, 0) + 1
        return breakdown


class IntakeAgent(BaseAgent):
    """
    Specialized agent for document intake, processing, and classification.

    Enhanced with comprehensive WCB case tracking:
    - Document source identification (insurance vs user vs medical providers)
    - Medical provider tracking (insurance doctors vs user doctors)
    - Key points extraction and tracking
    - Submission timeline management
    - Arguments and counter-arguments identification
    """

    def __init__(
        self,
        agent_id: Optional[str] = None,
        config: Optional[AgentConfig] = None,
        global_settings_manager=None,
        agent_coordinator=None,
        intake_directory: str = "./tmp/intake",
        processed_directory: str = "./tmp/processed",
        max_file_size_mb: int = 100,
        enable_ocr: bool = True,
        ocr_confidence_threshold: int = 70,
        use_mcp_filesystem: bool = True,
        enable_wcb_tracking: bool = True,
        **kwargs,
    ):
        super().__init__(
            agent_id=agent_id,
            config=config,
            global_settings_manager=global_settings_manager,
            **kwargs,
        )
        self.agent_coordinator = agent_coordinator

        # Intake-specific configuration
        self.intake_directory = Path(intake_directory)
        self.processed_directory = Path(processed_directory)
        self.max_file_size_mb = max_file_size_mb
        self.enable_ocr = enable_ocr and HAS_OCR
        self.ocr_confidence_threshold = ocr_confidence_threshold
        self.use_mcp_filesystem = use_mcp_filesystem
        self.enable_wcb_tracking = enable_wcb_tracking

        # Create directories if they don't exist
        self.intake_directory.mkdir(parents=True, exist_ok=True)
        self.processed_directory.mkdir(parents=True, exist_ok=True)

        # Document processing stats
        self.documents_processed = 0
        self.classification_cache: Dict[str, str] = {}

        # WCB case tracking
        self.wcb_trackers: Dict[str, WCBTracker] = {}
        self.current_case_tracker: Optional[WCBTracker] = None

        # Medical provider patterns for identification
        self.medical_title_patterns = [
            r"dr\.?\s+([a-z\s]+)",
            r"doctor\s+([a-z\s]+)",
            r"physician\s+([a-z\s]+)",
            r"([a-z\s]+),?\s+m\.?d\.?",
            r"([a-z\s]+),?\s+md",
            r"specialist\s+([a-z\s]+)",
            r"consultant\s+([a-z\s]+)",
        ]

        # Insurance company patterns
        self.insurance_company_patterns = [
            r"workers.*compensation.*board",
            r"wcb",
            r"worksafe",
            r"insurance.*company",
            r"claims.*adjuster",
            r"insurance.*adjuster",
            r"underwriter",
        ]

        # MCP and OCR status logging
        if self.use_mcp_filesystem:
            logger.info("🚀 Using MCP Rust filesystem for high-speed file operations")

        if self.enable_ocr:
            logger.info(
                f"✅ OCR enabled with Tesseract (confidence threshold: {ocr_confidence_threshold})"
            )
        else:
            if not HAS_OCR:
                logger.warning(
                    "⚠️ OCR disabled - missing dependencies (pytesseract, PIL, cv2)"
                )
            else:
                logger.info("ℹ️ OCR disabled by configuration")

        if self.enable_wcb_tracking:
            logger.info(
                "📊 WCB case tracking enabled - will sort insurance vs user submissions"
            )

        logger.info(f"IntakeAgent initialized: {self.agent_id}")

    async def _get_llm_instance(self):
        """Get LLM instance from the unified factory using global settings."""
        try:
            from ...utils.unified_llm_factory import get_llm_factory

            factory = get_llm_factory()

            # Use global settings manager if available
            if (
                hasattr(self, "global_settings_manager")
                and self.global_settings_manager
            ):
                # Get primary LLM configuration from global settings
                logger.info("🤖 Using LLM from global settings manager")
                return factory.create_llm_from_global_settings(
                    self.global_settings_manager, "primary"
                )
            else:
                # Fallback to default Google Gemini Pro if no global settings manager
                logger.warning(
                    "⚠️ No global settings manager found, using default Google Gemini Pro"
                )
                config = {
                    "provider": "google",
                    "model_name": "gemini-1.5-pro-latest",
                    "temperature": 0.7,
                    "max_tokens": 4096,
                }
                return factory.create_llm_from_config(config)

        except Exception as e:
            logger.warning(f"Could not get LLM instance: {e}")
            return None

    async def _mcp_read_file(self, file_path: Path) -> Dict[str, Any]:
        """Read file using MCP Rust filesystem for faster operations."""
        if not self.use_mcp_filesystem or not self.agent_coordinator:
            logger.info(
                f"MCP Filesystem disabled or coordinator not available. Reading {file_path.name} using standard Python I/O."
            )
            # Fallback to Python file operations
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                return {
                    "success": True,
                    "content": content,
                    "method": "python_filesystem_fallback",
                }
            except Exception as e:
                logger.warning(f"Standard file read failed for {file_path.name}: {e}")
                return {
                    "success": False,
                    "error": str(e),
                    "method": "python_filesystem_fallback",
                }

        try:
            logger.info(f"🚀 Using MCP Filesystem to read: {file_path.name}")
            mcp_result = await self.agent_coordinator.execute_agent_task(
                agent_type="filesystem_server",
                task_type="read_file",
                parameters={"file_path": str(file_path)},
            )
            if mcp_result.get("success"):
                return {
                    "success": True,
                    "content": mcp_result.get("content"),
                    "method": "mcp_filesystem",
                }
            else:
                logger.warning(
                    f"MCP file read failed for {file_path.name}: {mcp_result.get('error')}"
                )
                return {
                    "success": False,
                    "error": mcp_result.get("error"),
                    "method": "mcp_filesystem",
                }
        except Exception as e:
            logger.warning(f"MCP file read execution failed for {file_path.name}: {e}")
            return {
                "success": False,
                "error": str(e),
                "method": "mcp_filesystem_exception",
            }

    async def _mcp_move_file(
        self, source_path: Path, target_path: Path
    ) -> Dict[str, Any]:
        """Move file using MCP Rust filesystem for faster operations."""
        if not self.use_mcp_filesystem or not self.agent_coordinator:
            logger.info(
                f"MCP Filesystem disabled or coordinator not available. Moving {source_path.name} using standard Python I/O."
            )
            # Fallback to Python file operations
            try:
                import shutil

                shutil.move(str(source_path), str(target_path))
                return {
                    "success": True,
                    "source": str(source_path),
                    "destination": str(target_path),
                    "method": "python_filesystem_fallback",
                }
            except Exception as e:
                logger.warning(f"Standard file move failed for {source_path.name}: {e}")
                return {
                    "success": False,
                    "error": str(e),
                    "method": "python_filesystem_fallback",
                }

        try:
            logger.info(
                f"🚀 Using MCP Filesystem to move: {source_path.name} to {target_path}"
            )
            # Assuming the rust-mcp-filesystem.exe has a 'move_file' tool
            # And it takes 'source_path' and 'target_path' as parameters
            mcp_result = await self.agent_coordinator.execute_agent_task(
                agent_type="filesystem_server",
                task_type="move_file",  # Ensure this tool exists and params match
                parameters={
                    "source_path": str(source_path),
                    "destination_path": str(target_path),
                    "overwrite": True,
                },
            )
            if mcp_result.get("success"):
                return {
                    "success": True,
                    "source": str(source_path),
                    "destination": str(
                        target_path
                    ),  # Assuming mcp_result contains this
                    "method": "mcp_filesystem",
                }
            else:
                logger.warning(
                    f"MCP file move failed for {source_path.name}: {mcp_result.get('error')}"
                )
                return {
                    "success": False,
                    "error": mcp_result.get("error"),
                    "method": "mcp_filesystem",
                }
        except Exception as e:
            logger.warning(
                f"MCP file move execution failed for {source_path.name}: {e}"
            )
            return {
                "success": False,
                "error": str(e),
                "method": "mcp_filesystem_exception",
            }

    async def _mcp_get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get file information using MCP Rust filesystem."""
        if not self.use_mcp_filesystem or not self.agent_coordinator:
            logger.info(
                f"MCP Filesystem disabled or coordinator not available. Getting info for {file_path.name} using standard Python I/O."
            )
            # Fallback to Python stat
            try:
                stat = file_path.stat()
                return {
                    "success": True,
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "method": "python_filesystem_fallback",
                }
            except Exception as e:
                logger.warning(f"Standard file info failed for {file_path.name}: {e}")
                return {
                    "success": False,
                    "error": str(e),
                    "method": "python_filesystem_fallback",
                }

        try:
            logger.info(f"🚀 Using MCP Filesystem to get info for: {file_path.name}")
            mcp_result = await self.agent_coordinator.execute_agent_task(
                agent_type="filesystem_server",
                task_type="get_file_info",  # Ensure this tool exists
                parameters={"file_path": str(file_path)},
            )
            if mcp_result.get("success"):
                # Assuming mcp_result contains 'size' and 'modified' or similar
                return {
                    "success": True,
                    "size": mcp_result.get("size"),
                    "modified": mcp_result.get(
                        "modified_time"
                    ),  # Adjust key based on actual MCP tool output
                    "method": "mcp_filesystem",
                }
            else:
                logger.warning(
                    f"MCP file info failed for {file_path.name}: {mcp_result.get('error')}"
                )
                return {
                    "success": False,
                    "error": mcp_result.get("error"),
                    "method": "mcp_filesystem",
                }
        except Exception as e:
            logger.warning(f"MCP file info execution failed for {file_path.name}: {e}")
            return {
                "success": False,
                "error": str(e),
                "method": "mcp_filesystem_exception",
            }

    async def _mcp_list_directory(self, directory_path: Path) -> Dict[str, Any]:
        """List directory contents using MCP Rust filesystem."""
        if not self.use_mcp_filesystem or not self.agent_coordinator:
            logger.info(
                f"MCP Filesystem disabled or coordinator not available. Listing {directory_path} using standard Python I/O."
            )
            # Fallback to Python listdir
            try:
                files = list(directory_path.iterdir())
                return {
                    "success": True,
                    "files": [str(f) for f in files],
                    "count": len(files),
                    "method": "python_filesystem_fallback",
                }
            except Exception as e:
                logger.warning(
                    f"Standard directory list failed for {directory_path}: {e}"
                )
                return {
                    "success": False,
                    "error": str(e),
                    "method": "python_filesystem_fallback",
                }

        try:
            logger.info(f"🚀 Using MCP Filesystem to list directory: {directory_path}")
            mcp_result = await self.agent_coordinator.execute_agent_task(
                agent_type="filesystem_server",
                task_type="list_directory",  # Ensure this tool exists
                parameters={"path": str(directory_path)},  # Adjust param name if needed
            )
            if mcp_result.get("success"):
                # Assuming mcp_result contains 'files' (list of strings) and 'count'
                return {
                    "success": True,
                    "files": mcp_result.get("files"),
                    "count": mcp_result.get("count"),
                    "method": "mcp_filesystem",
                }
            else:
                logger.warning(
                    f"MCP directory list failed for {directory_path}: {mcp_result.get('error')}"
                )
                return {
                    "success": False,
                    "error": mcp_result.get("error"),
                    "method": "mcp_filesystem",
                }
        except Exception as e:
            logger.warning(
                f"MCP directory list execution failed for {directory_path}: {e}"
            )
            return {
                "success": False,
                "error": str(e),
                "method": "mcp_filesystem_exception",
            }

    def get_supported_task_types(self) -> List[str]:
        """Return list of task types this agent can handle."""
        base_tasks = [
            "process_document",
            "classify_document",
            "extract_metadata",
            "validate_file",
            "batch_process",
            "batch_process_optimized",  # Optimized batch processing (OCR last)
            "similarity_check",
        ]

        if self.enable_ocr:
            base_tasks.extend(
                [
                    "ocr_extract",
                ]
            )

        if self.enable_wcb_tracking:
            base_tasks.extend(
                [
                    "start_case_tracking",
                    "get_case_tracking_summary",
                    "identify_document_source",
                    "extract_medical_providers",
                    "extract_key_points",
                ]
            )

        return base_tasks

    async def execute_task(self, task: AgentTask) -> Dict[str, Any]:
        """Execute a specific intake task."""
        task_type = task.task_type
        params = task.parameters

        try:
            if task_type == "process_document":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._process_document(
                    file_path=file_path,
                    document_id=params.get("document_id"),
                    metadata=params.get("metadata", {}),
                )

            elif task_type == "classify_document":
                content = params.get("content")
                if not content:
                    raise ValueError("content parameter is required")
                return await self._classify_document(
                    content=content, file_path=params.get("file_path")
                )

            elif task_type == "extract_metadata":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._extract_metadata(file_path=Path(file_path))

            elif task_type == "validate_file":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._validate_file(file_path=file_path)

            elif task_type == "batch_process":
                return await self._batch_process(
                    directory=params.get("directory", self.intake_directory),
                    file_patterns=params.get(
                        "file_patterns",
                        ["*.pdf", "*.docx", "*.txt", "*.png", "*.jpg", "*.jpeg"],
                    ),
                )

            elif task_type == "batch_process_optimized":
                return await self._batch_process_optimized(
                    directory=params.get("directory", self.intake_directory),
                    file_patterns=params.get(
                        "file_patterns",
                        ["*.pdf", "*.docx", "*.txt", "*.png", "*.jpg", "*.jpeg"],
                    ),
                )

            elif task_type == "ocr_extract":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._ocr_extract(file_path=file_path)

            elif task_type == "similarity_check":
                content = params.get("content")
                document_id = params.get("document_id")
                if not content or not document_id:
                    raise ValueError("content and document_id parameters are required")
                return await self._similarity_check(
                    content=content, document_id=document_id
                )

            elif task_type == "start_case_tracking":
                case_id = params.get("case_id")
                if not case_id:
                    raise ValueError("case_id parameter is required")
                return await self._start_case_tracking(case_id)

            elif task_type == "get_case_tracking_summary":
                case_id = params.get("case_id")
                if not case_id:
                    raise ValueError("case_id parameter is required")
                return await self._get_case_tracking_summary(case_id)

            elif task_type == "identify_document_source":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._identify_document_source(file_path)

            elif task_type == "extract_medical_providers":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._extract_medical_providers(file_path)

            elif task_type == "extract_key_points":
                file_path = params.get("file_path")
                if not file_path:
                    raise ValueError("file_path parameter is required")
                return await self._extract_key_points(file_path)

            else:
                raise ValueError(f"Unsupported task type: {task_type}")

        except Exception as e:
            logger.error(f"Task execution failed for {task_type}: {e}", exc_info=True)
            raise

    async def _process_document(
        self,
        file_path: Union[str, Path],
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Process a single document through the full intake pipeline."""
        file_path_obj = Path(file_path)
        document_id = document_id or str(uuid.uuid4())

        logger.info(f"Processing document: {file_path_obj}")

        # Step 1: Validate file
        validation_result = await self._validate_file(file_path_obj)
        if not validation_result["valid"]:
            return {
                "success": False,
                "document_id": document_id,
                "error": validation_result["error"],
                "stage": "validation",
            }

        # Step 2: Extract text content
        content_result = await self._extract_text_content(file_path_obj)
        if not content_result["success"]:
            return {
                "success": False,
                "document_id": document_id,
                "error": content_result["error"],
                "stage": "content_extraction",
            }

        # Step 3: Extract metadata
        metadata_result = await self._extract_metadata(file_path_obj)

        # Step 4: Classify document
        classification_result = await self._classify_document(
            content=content_result["content"], file_path=file_path_obj
        )

        # Step 5: Perform similarity check if requested
        similarity_result = None
        if len(content_result["content"]) > 100:  # Only for substantial content
            similarity_result = await self._similarity_check(
                content=content_result["content"], document_id=document_id
            )

        # Compile final result
        result_data = {
            "success": True,
            "document_id": document_id,
            "file_path": str(file_path_obj),
            "document_type": validation_result["document_type"],
            "file_size_mb": validation_result["file_size_mb"],
            "content": content_result["content"],
            "content_length": len(content_result["content"]),
            "metadata": {
                **metadata_result.get("metadata", {}),
                **(metadata or {}),
                "processed_at": datetime.now().isoformat(),
                "processed_by": self.agent_id,
            },
            "classification": classification_result,
            "similarity": similarity_result,
        }

        # Move file to processed directory
        processed_path = (
            self.processed_directory / f"{document_id}_{file_path_obj.name}"
        )
        try:
            if file_path_obj != processed_path:
                # Use _mcp_move_file or a copy then delete approach if move is not robust
                # For simplicity, let's try to read content and write to new path via MCP
                # This assumes the file isn't too large for memory.
                # A true 'copy_file' MCP tool would be better.
                if self.use_mcp_filesystem and self.agent_coordinator:
                    # Read content first (already done in content_result)
                    original_content = content_result[
                        "content"
                    ]  # Or re-read if binary/large

                    # If original file was binary or needs byte-perfect copy, re-read as bytes
                    # This is a simplification; binary files need different handling for 'content'
                    # For now, assuming text-based or OCRed content is what we're "moving"

                    write_mcp_result = await self.agent_coordinator.execute_agent_task(
                        agent_type="filesystem_server",
                        task_type="write_file",
                        parameters={
                            "file_path": str(processed_path),
                            "content": original_content,
                            "overwrite": True,
                        },
                    )
                    if write_mcp_result.get("success"):
                        result_data["processed_path"] = str(processed_path)
                        # Optionally, delete original if it's a true move, but be careful
                        # delete_mcp_result = await self.agent_coordinator.execute_agent_task(
                        # agent_type="filesystem_server",
                        # task_type="delete_file",
                        # parameters={"file_path": str(file_path_obj)}
                        # )
                        # if not delete_mcp_result.get("success"):
                        # logger.warning(f"MCP failed to delete original file {file_path_obj.name} after copy: {delete_mcp_result.get('error')}")
                    else:
                        logger.warning(
                            f"MCP failed to write processed file {processed_path.name}: {write_mcp_result.get('error')}"
                        )
                        # Fallback to shutil if MCP write fails? Or just log error?
                        # For now, just log.
                else:
                    import shutil

                    shutil.copy2(file_path_obj, processed_path)  # Fallback to copy
                    result_data["processed_path"] = str(processed_path)
        except Exception as e:
            logger.warning(f"Failed to move file to processed directory: {e}")

        self.documents_processed += 1
        logger.info(f"Document processed successfully: {document_id}")

        return result_data

    async def _validate_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Validate file format, size, and accessibility."""
        file_path = Path(file_path)

        try:
            if not file_path.exists():
                return {
                    "valid": False,
                    "error": "File does not exist",
                    "file_path": str(file_path),
                }

            if not file_path.is_file():
                return {
                    "valid": False,
                    "error": "Path is not a file",
                    "file_path": str(file_path),
                }

            # Check file size
            file_size_bytes = file_path.stat().st_size
            file_size_mb = file_size_bytes / (1024 * 1024)

            if file_size_mb > self.max_file_size_mb:
                return {
                    "valid": False,
                    "error": f"File too large: {file_size_mb:.2f}MB > {self.max_file_size_mb}MB",
                    "file_size_mb": file_size_mb,
                }

            # Determine document type
            document_type = self._determine_document_type(file_path)

            return {
                "valid": True,
                "file_path": str(file_path),
                "file_size_mb": file_size_mb,
                "document_type": document_type,
            }

        except Exception as e:
            return {
                "valid": False,
                "error": f"File validation error: {str(e)}",
                "file_path": str(file_path),
            }

    def _determine_document_type(self, file_path: Path) -> str:
        """Determine document type based on file extension and content."""
        extension = file_path.suffix.lower()

        type_mapping = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".json": DocumentType.JSON,
            ".png": DocumentType.IMAGE,
            ".jpg": DocumentType.IMAGE,
            ".jpeg": DocumentType.IMAGE,
            ".tiff": DocumentType.IMAGE,
            ".tif": DocumentType.IMAGE,
            ".bmp": DocumentType.IMAGE,
            ".gif": DocumentType.IMAGE,
        }

        return type_mapping.get(extension, DocumentType.UNKNOWN)

    def _is_ocr_required(self, file_path: Path) -> bool:
        """Check if file requires OCR processing."""
        document_type = self._determine_document_type(file_path)
        return document_type in [DocumentType.IMAGE, DocumentType.SCANNED_PDF]

    async def _extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text content from various file formats."""
        document_type = self._determine_document_type(file_path)

        try:
            if document_type == DocumentType.PDF:
                # Try regular PDF extraction first
                pdf_result = await self._extract_pdf_content(file_path)

                # If PDF extraction fails or returns minimal content, try OCR
                if (
                    not pdf_result["success"]
                    or len(pdf_result.get("content", "").strip()) < 50
                ) and self.enable_ocr:
                    logger.info(
                        f"📄 PDF appears to be scanned, attempting OCR: {file_path.name}"
                    )
                    ocr_result = await self._ocr_extract_pdf(file_path)
                    if ocr_result["success"]:
                        return ocr_result

                return pdf_result

            elif document_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            elif document_type == DocumentType.TXT:
                return await self._extract_txt_content(file_path)
            elif document_type == DocumentType.HTML:
                return await self._extract_html_content(file_path)
            elif document_type == DocumentType.JSON:
                return await self._extract_json_content(file_path)
            elif document_type == DocumentType.IMAGE and self.enable_ocr:
                return await self._ocr_extract_image(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document type: {document_type}",
                    "content": "",
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"Content extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files."""
        if not HAS_PYPDF2:
            return {
                "success": False,
                "error": "PyPDF2 not available. Install with: pip install PyPDF2",
                "content": "",
            }

        try:
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                content = ""

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )

                return {
                    "success": True,
                    "content": content.strip(),
                    "page_count": len(reader.pages),
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        if not HAS_PYTHON_DOCX:
            return {
                "success": False,
                "error": "python-docx not available. Install with: pip install python-docx",
                "content": "",
            }

        try:
            doc = docx.Document(str(file_path))  # Convert Path to string
            content = ""

            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            return {
                "success": True,
                "content": content.strip(),
                "paragraph_count": len(doc.paragraphs),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_txt_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()

            return {
                "success": True,
                "content": content,
                "character_count": len(content),
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_html_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from HTML files."""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                soup = BeautifulSoup(file.read(), "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text()

            return {
                "success": True,
                "content": content.strip(),
                "html_title": soup.title.string if soup.title else None,
            }
        except ImportError:
            return {
                "success": False,
                "error": "BeautifulSoup not available. Install with: pip install beautifulsoup4",
                "content": "",
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"HTML extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_json_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from JSON files."""
        try:
            import json

            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Convert JSON to readable text
            content = json.dumps(data, indent=2, ensure_ascii=False)

            return {
                "success": True,
                "content": content,
                "json_keys": list(data.keys()) if isinstance(data, dict) else None,
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"JSON extraction failed: {str(e)}",
                "content": "",
            }

    async def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file."""
        try:
            stat = file_path.stat()

            metadata = {
                "file_name": file_path.name,
                "file_extension": file_path.suffix,
                "file_size_bytes": stat.st_size,
                "file_size_mb": stat.st_size / (1024 * 1024),
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_time": datetime.fromtimestamp(stat.st_atime).isoformat(),
            }

            return {"success": True, "metadata": metadata}

        except Exception as e:
            return {
                "success": False,
                "error": f"Metadata extraction failed: {str(e)}",
                "metadata": {},
            }

    async def _classify_document(
        self, content: str, file_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Classify document type and content category using LLM analysis."""
        if not content or len(content.strip()) < 50:
            return {
                "case_type": CaseType.UNKNOWN,
                "confidence": 0.0,
                "reasoning": "Insufficient content for classification",
            }

        try:
            # Use LLM for intelligent classification
            classification_prompt = f"""
            Analyze the following document content and classify it into one of these categories:
            1. legal_case - Legal documents, court cases, legal briefs
            2. research_paper - Academic papers, research studies, scientific documents
            3. technical_doc - Technical documentation, manuals, specifications
            4. business_doc - Business reports, proposals, correspondence
            5. personal_doc - Personal documents, letters, notes
            6. unknown - Cannot be classified

            Content to analyze:
            {content[:2000]}...

            Respond with JSON format:
            {{
                "case_type": "category_name",
                "confidence": 0.0-1.0,
                "reasoning": "Brief explanation",
                "key_indicators": ["list", "of", "indicators"]
            }}
            """

            response = await self.llm.ainvoke(classification_prompt)

            # Parse LLM response
            try:
                import json

                result = json.loads(response.content)

                # Validate classification
                valid_types = [
                    CaseType.LEGAL_CASE,
                    CaseType.RESEARCH_PAPER,
                    CaseType.TECHNICAL_DOC,
                    CaseType.BUSINESS_DOC,
                    CaseType.PERSONAL_DOC,
                    CaseType.UNKNOWN,
                ]

                if result.get("case_type") not in valid_types:
                    result["case_type"] = CaseType.UNKNOWN
                    result["confidence"] = 0.0

                return result

            except (json.JSONDecodeError, KeyError):
                # Fallback to simple keyword-based classification
                return self._fallback_classify(content, file_path)

        except Exception as e:
            logger.warning(f"LLM classification failed: {e}")
            return self._fallback_classify(content, file_path)

    def _fallback_classify(
        self, content: str, file_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """Fallback classification using simple keyword matching."""
        content_lower = content.lower()

        # Legal indicators
        legal_keywords = [
            "court",
            "plaintiff",
            "defendant",
            "judgment",
            "appeal",
            "statute",
            "legal",
            "law",
            "case",
            "trial",
            "hearing",
            "docket",
        ]
        legal_score = sum(1 for keyword in legal_keywords if keyword in content_lower)

        # Research indicators
        research_keywords = [
            "abstract",
            "methodology",
            "conclusion",
            "references",
            "study",
            "research",
            "analysis",
            "experiment",
            "hypothesis",
            "data",
        ]
        research_score = sum(
            1 for keyword in research_keywords if keyword in content_lower
        )

        # Technical indicators
        tech_keywords = [
            "specification",
            "manual",
            "documentation",
            "technical",
            "system",
            "configuration",
            "implementation",
            "architecture",
            "api",
            "protocol",
        ]
        tech_score = sum(1 for keyword in tech_keywords if keyword in content_lower)

        # Business indicators
        business_keywords = [
            "proposal",
            "report",
            "meeting",
            "project",
            "budget",
            "revenue",
            "business",
            "management",
            "strategy",
            "market",
            "client",
        ]
        business_score = sum(
            1 for keyword in business_keywords if keyword in content_lower
        )

        # Determine classification based on highest score
        scores = {
            CaseType.LEGAL_CASE: legal_score,
            CaseType.RESEARCH_PAPER: research_score,
            CaseType.TECHNICAL_DOC: tech_score,
            CaseType.BUSINESS_DOC: business_score,
        }

        max_score = max(scores.values())
        if max_score == 0:
            return {
                "case_type": CaseType.UNKNOWN,
                "confidence": 0.0,
                "reasoning": "No clear indicators found",
                "key_indicators": [],
            }

        case_type = max(scores, key=lambda k: scores[k])
        confidence = min(max_score / 10.0, 1.0)  # Normalize to 0-1

        return {
            "case_type": case_type,
            "confidence": confidence,
            "reasoning": f"Keyword-based classification (score: {max_score})",
            "key_indicators": [k for k, v in scores.items() if v > 0],
        }

    async def _similarity_check(self, content: str, document_id: str) -> Dict[str, Any]:
        """Perform basic similarity analysis with existing documents."""
        # This is a placeholder for similarity analysis
        # In a full implementation, this would compare against a database
        # of existing documents using embedding similarity

        try:
            # Simple similarity check based on content length and structure
            word_count = len(content.split())
            unique_words = len(set(content.lower().split()))

            similarity_metrics = {
                "word_count": word_count,
                "unique_words": unique_words,
                "vocabulary_richness": unique_words / max(word_count, 1),
                "similar_documents": [],  # Would be populated from database search
                "similarity_score": 0.0,  # Would be calculated against existing docs
            }

            return {
                "success": True,
                "document_id": document_id,
                "metrics": similarity_metrics,
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Similarity check failed: {str(e)}",
                "document_id": document_id,
            }

    async def _batch_process(
        self, directory: Union[str, Path], file_patterns: List[str]
    ) -> Dict[str, Any]:
        """Process multiple documents in a directory."""
        directory = Path(directory)

        if not directory.exists():
            return {
                "success": False,
                "error": f"Directory does not exist: {directory}",
                "processed_count": 0,
            }

        results = []
        processed_count = 0
        error_count = 0

        try:
            for pattern in file_patterns:
                for file_path in directory.glob(pattern):
                    try:
                        result = await self._process_document(file_path)
                        results.append(result)

                        if result["success"]:
                            processed_count += 1
                        else:
                            error_count += 1

                    except Exception as e:
                        error_count += 1
                        results.append(
                            {
                                "success": False,
                                "file_path": str(file_path),
                                "error": str(e),
                            }
                        )

            return {
                "success": True,
                "directory": str(directory),
                "processed_count": processed_count,
                "error_count": error_count,
                "total_files": len(results),
                "results": results,
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Batch processing failed: {str(e)}",
                "processed_count": processed_count,
                "error_count": error_count,
            }

    async def get_intake_statistics(self) -> Dict[str, Any]:
        """Get statistics about document intake operations."""
        return {
            "agent_id": self.agent_id,
            "documents_processed": self.documents_processed,
            "intake_directory": str(self.intake_directory),
            "processed_directory": str(self.processed_directory),
            "max_file_size_mb": self.max_file_size_mb,
            "supported_formats": [
                DocumentType.PDF,
                DocumentType.DOCX,
                DocumentType.TXT,
                DocumentType.HTML,
                DocumentType.JSON,
            ],
            "cache_size": len(self.classification_cache),
        }

    async def _ocr_extract(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract text using OCR from images or scanned PDFs."""
        if not self.enable_ocr:
            return {
                "success": False,
                "error": "OCR is not enabled or dependencies are missing",
                "content": "",
            }

        file_path = Path(file_path)
        document_type = self._determine_document_type(file_path)

        try:
            if document_type == DocumentType.IMAGE:
                return await self._ocr_extract_image(file_path)
            elif document_type == DocumentType.PDF:
                return await self._ocr_extract_pdf(file_path)
            else:
                return {
                    "success": False,
                    "error": f"OCR not supported for document type: {document_type}",
                    "content": "",
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"OCR extraction failed: {str(e)}",
                "content": "",
            }

    async def _ocr_extract_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from image files using OCR."""
        try:
            # Load and preprocess image
            image = cv2.imread(str(file_path))
            if image is None:
                # Try with PIL if cv2 fails
                pil_image = Image.open(file_path)
                image = np.array(pil_image)

            # Preprocess image for better OCR results
            processed_image = self._preprocess_image_for_ocr(image)

            # Extract text using pytesseract
            custom_config = '--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,!?;:()[]{}"-'

            # Get detailed data for confidence checking
            data = pytesseract.image_to_data(
                processed_image,
                output_type=pytesseract.Output.DICT,
                config=custom_config,
            )

            # Filter text by confidence
            filtered_text = []
            for i, conf in enumerate(data["conf"]):
                if int(conf) > self.ocr_confidence_threshold:
                    text = data["text"][i].strip()
                    if text:
                        filtered_text.append(text)

            content = " ".join(filtered_text)

            # Also get raw text as fallback
            raw_text = pytesseract.image_to_string(
                processed_image, config=custom_config
            )

            # Use filtered text if it's substantial, otherwise use raw text
            final_content = content if len(content) > len(raw_text) * 0.7 else raw_text

            return {
                "success": True,
                "content": final_content.strip(),
                "method": "ocr_image",
                "confidence_threshold": self.ocr_confidence_threshold,
                "raw_text_length": len(raw_text),
                "filtered_text_length": len(content),
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Image OCR failed: {str(e)}",
                "content": "",
            }

    async def _ocr_extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using OCR (for scanned PDFs)."""
        try:
            # Convert PDF pages to images using pdf2image or similar
            # For now, we'll use a simple approach with PIL
            logger.info(f"🔍 Attempting OCR on PDF: {file_path.name}")

            # This is a simplified implementation
            # In a full implementation, you'd convert PDF pages to images first
            content = f"OCR extraction attempted for PDF: {file_path.name}\n"
            content += "Note: Full PDF-to-OCR conversion requires additional libraries like pdf2image"

            return {
                "success": True,
                "content": content,
                "method": "ocr_pdf_placeholder",
                "note": "Full PDF OCR requires pdf2image library integration",
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"PDF OCR failed: {str(e)}",
                "content": "",
            }

    def _preprocess_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image to improve OCR accuracy."""
        try:
            # Convert to grayscale
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image

            # Normalize the image
            norm_img = np.zeros((gray.shape[0], gray.shape[1]))
            gray = cv2.normalize(gray, norm_img, 0, 255, cv2.NORM_MINMAX)

            # Apply threshold to get image with only black and white
            gray = cv2.threshold(gray, 100, 255, cv2.THRESH_BINARY)[1]

            # Apply slight Gaussian blur to reduce noise
            gray = cv2.GaussianBlur(gray, (1, 1), 0)

            return gray

        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image

    async def _batch_process_optimized(
        self, directory: Union[str, Path], file_patterns: List[str]
    ) -> Dict[str, Any]:
        """
        Optimized batch processing that processes OCR documents last for better speed.

        Processing order:
        1. Text files (.txt, .json)
        2. Document files (.docx, .pdf with text)
        3. OCR files (.png, .jpg, scanned PDFs) - processed last
        """
        directory = Path(directory)

        if not directory.exists():
            return {
                "success": False,
                "error": f"Directory does not exist: {directory}",
                "processed_count": 0,
            }

        logger.info(f"🚀 Starting optimized batch processing: {directory}")

        # Categorize files by processing priority
        files_by_priority = {
            "fast": [],  # Text files, JSON
            "medium": [],  # DOCX, text-based PDFs
            "slow": [],  # Images, scanned PDFs (OCR required)
        }

        # Scan and categorize all files
        for pattern in file_patterns:
            for file_path in directory.glob(pattern):
                if file_path.is_file():
                    if self._is_ocr_required(file_path):
                        files_by_priority["slow"].append(file_path)
                    elif file_path.suffix.lower() in {".txt", ".json", ".html"}:
                        files_by_priority["fast"].append(file_path)
                    else:
                        files_by_priority["medium"].append(file_path)

        logger.info("📊 File categorization:")
        logger.info(f"   🚄 Fast files (text/json): {len(files_by_priority['fast'])}")
        logger.info(
            f"   🚗 Medium files (docs/pdfs): {len(files_by_priority['medium'])}"
        )
        logger.info(
            f"   🐌 Slow files (OCR required): {len(files_by_priority['slow'])}"
        )

        results = []
        processed_count = 0
        error_count = 0
        start_time = datetime.now()

        # Process in order of speed: fast -> medium -> slow
        for priority in ["fast", "medium", "slow"]:
            priority_files = files_by_priority[priority]

            if priority_files:
                logger.info(
                    f"🔄 Processing {priority} priority files ({len(priority_files)} files)..."
                )

                for file_path in priority_files:
                    try:
                        result = await self._process_document(file_path)
                        results.append(result)

                        if result["success"]:
                            processed_count += 1
                            logger.info(f"✅ Processed: {file_path.name}")
                        else:
                            error_count += 1
                            logger.warning(
                                f"❌ Failed: {file_path.name} - {result.get('error')}"
                            )

                    except Exception as e:
                        error_count += 1
                        error_result = {
                            "success": False,
                            "file_path": str(file_path),
                            "error": str(e),
                        }
                        results.append(error_result)
                        logger.error(f"💥 Exception processing {file_path.name}: {e}")

                logger.info(f"✅ Completed {priority} priority files")

        end_time = datetime.now()
        processing_duration = (end_time - start_time).total_seconds()

        return {
            "success": True,
            "directory": str(directory),
            "processed_count": processed_count,
            "error_count": error_count,
            "total_files": len(results),
            "processing_duration_seconds": processing_duration,
            "optimization_stats": {
                "fast_files": len(files_by_priority["fast"]),
                "medium_files": len(files_by_priority["medium"]),
                "slow_files_ocr": len(files_by_priority["slow"]),
            },
            "results": results,
        }

    async def _start_case_tracking(self, case_id: str) -> Dict[str, Any]:
        """Start tracking a new case."""
        if case_id in self.wcb_trackers:
            return {
                "success": False,
                "error": f"Case tracking already started for case: {case_id}",
            }

        self.wcb_trackers[case_id] = WCBTracker(case_id)
        self.current_case_tracker = self.wcb_trackers[case_id]
        return {
            "success": True,
            "case_id": case_id,
        }

    async def _get_case_tracking_summary(self, case_id: str) -> Dict[str, Any]:
        """Get comprehensive case tracking summary."""
        if case_id not in self.wcb_trackers:
            return {
                "success": False,
                "error": f"Case tracking not started for case: {case_id}",
            }

        summary = self.wcb_trackers[case_id].get_summary()
        return {
            "success": True,
            "case_id": case_id,
            "summary": summary,
        }

    async def _identify_document_source(
        self, file_path: Union[str, Path]
    ) -> Dict[str, Any]:
        """Identify the source of a document (insurance company, user, medical provider, etc.)."""
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                "success": False,
                "error": "File does not exist",
                "file_path": str(file_path),
            }

        try:
            # Extract text content for analysis
            content_result = await self._extract_text_content(file_path)
            if not content_result["success"]:
                return {
                    "success": False,
                    "error": f"Could not extract text: {content_result['error']}",
                    "file_path": str(file_path),
                }

            content = content_result["content"].lower()
            filename = file_path.name.lower()

            # Analyze document source based on content and filename patterns
            source_scores = {
                DocumentSource.INSURANCE_COMPANY: 0,
                DocumentSource.USER_CLAIMANT: 0,
                DocumentSource.WCB_BOARD: 0,
                DocumentSource.MEDICAL_PROVIDER: 0,
                DocumentSource.EMPLOYER: 0,
                DocumentSource.LEGAL_REPRESENTATIVE: 0,
            }

            # Insurance company indicators
            insurance_patterns = [
                r"workers.*compensation.*board",
                r"wcb.*claim",
                r"insurance.*company",
                r"claims.*adjuster",
                r"underwriter",
                r"denying.*claim",
                r"claim.*denial",
                r"coverage.*determination",
                r"insurance.*policy",
                r"liability.*determination",
            ]

            for pattern in insurance_patterns:
                if re.search(pattern, content) or re.search(pattern, filename):
                    source_scores[DocumentSource.INSURANCE_COMPANY] += 1

            # User/claimant indicators
            user_patterns = [
                r"claimant.*statement",
                r"my.*injury",
                r"my.*claim",
                r"injured.*worker",
                r"personal.*statement",
                r"i.*was.*injured",
                r"my.*doctor",
                r"my.*treatment",
                r"appealing.*decision",
                r"request.*for.*reconsideration",
            ]

            for pattern in user_patterns:
                if re.search(pattern, content):
                    source_scores[DocumentSource.USER_CLAIMANT] += 1

            # Medical provider indicators
            medical_patterns = [
                r"medical.*report",
                r"physician.*report",
                r"doctor.*report",
                r"medical.*assessment",
                r"clinical.*notes",
                r"examination.*findings",
                r"diagnosis.*and.*treatment",
                r"medical.*opinion",
                r"independent.*medical.*exam",
                r"ime.*report",
            ]

            for pattern in medical_patterns:
                if re.search(pattern, content) or re.search(pattern, filename):
                    source_scores[DocumentSource.MEDICAL_PROVIDER] += 1

            # WCB Board indicators
            wcb_patterns = [
                r"workers.*compensation.*board",
                r"wcb.*decision",
                r"board.*decision",
                r"tribunal.*decision",
                r"appeal.*tribunal",
                r"wcat.*decision",
                r"reconsideration.*decision",
            ]

            for pattern in wcb_patterns:
                if re.search(pattern, content) or re.search(pattern, filename):
                    source_scores[DocumentSource.WCB_BOARD] += 1

            # Employer indicators
            employer_patterns = [
                r"employer.*report",
                r"incident.*report",
                r"workplace.*injury.*report",
                r"return.*to.*work",
                r"job.*description",
                r"modified.*duties",
                r"human.*resources",
            ]

            for pattern in employer_patterns:
                if re.search(pattern, content) or re.search(pattern, filename):
                    source_scores[DocumentSource.EMPLOYER] += 1

            # Legal representative indicators
            legal_patterns = [
                r"law.*firm",
                r"legal.*counsel",
                r"attorney",
                r"lawyer",
                r"legal.*representative",
                r"on.*behalf.*of",
                r"legal.*opinion",
                r"brief.*submitted",
            ]

            for pattern in legal_patterns:
                if re.search(pattern, content) or re.search(pattern, filename):
                    source_scores[DocumentSource.LEGAL_REPRESENTATIVE] += 1

            # Determine most likely source
            if max(source_scores.values()) == 0:
                identified_source = DocumentSource.UNKNOWN
                confidence = 0.0
            else:
                identified_source = max(source_scores, key=lambda k: source_scores[k])
                confidence = source_scores[identified_source] / sum(
                    source_scores.values()
                )

            return {
                "success": True,
                "file_path": str(file_path),
                "source": identified_source,
                "confidence": confidence,
                "source_scores": source_scores,
                "method": "pattern_analysis",
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Document source identification failed: {str(e)}",
                "file_path": str(file_path),
            }

    async def _extract_medical_providers(
        self, file_path: Union[str, Path]
    ) -> Dict[str, Any]:
        """Extract medical providers mentioned in a document."""
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                "success": False,
                "error": "File does not exist",
                "file_path": str(file_path),
            }

        try:
            # Extract text content
            content_result = await self._extract_text_content(file_path)
            if not content_result["success"]:
                return {
                    "success": False,
                    "error": f"Could not extract text: {content_result['error']}",
                    "file_path": str(file_path),
                }

            content = content_result["content"]
            providers_found = []
            specialties_found = []

            # Extract doctor names using various patterns
            for pattern in self.medical_title_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        # For patterns with multiple groups, take the name group
                        provider_name = (
                            match[0] if match[0] else match[1] if len(match) > 1 else ""
                        )
                    else:
                        provider_name = match

                    # Clean up the name
                    provider_name = provider_name.strip().title()
                    if len(provider_name) > 2 and provider_name not in providers_found:
                        providers_found.append(provider_name)

            # Extract medical specialties
            specialty_patterns = [
                r"orthopedic.*surgeon",
                r"neurologist",
                r"psychiatrist",
                r"psychologist",
                r"physiotherapist",
                r"occupational.*therapist",
                r"chiropractor",
                r"family.*physician",
                r"general.*practitioner",
                r"specialist.*in.*([a-z\s]+)",
                r"([a-z\s]+).*specialist",
            ]

            for pattern in specialty_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple):
                        specialty = (
                            match[0] if match[0] else match[1] if len(match) > 1 else ""
                        )
                    else:
                        specialty = match

                    specialty = specialty.strip().title()
                    if len(specialty) > 2 and specialty not in specialties_found:
                        specialties_found.append(specialty)

            # Use LLM for more sophisticated extraction if available
            llm_providers = []
            try:
                llm = await self._get_llm_instance()
                if llm:
                    prompt = f"""
                    Analyze this medical document and extract all medical providers mentioned.
                    Include doctor names, clinics, hospitals, and medical specialists.

                    Document content: {content[:3000]}...

                    Return only a JSON list of medical providers found:
                    ["Dr. Smith", "ABC Medical Clinic", "City Hospital Orthopedic Department"]
                    """

                    response = await llm.ainvoke(prompt)
                    try:
                        import json

                        llm_providers = json.loads(response.content)
                        if not isinstance(llm_providers, list):
                            llm_providers = []
                    except:
                        llm_providers = []
            except Exception as e:
                logger.warning(f"LLM medical provider extraction failed: {e}")

            # Combine pattern-based and LLM results
            all_providers = list(set(providers_found + llm_providers))

            return {
                "success": True,
                "file_path": str(file_path),
                "providers": all_providers,
                "specialties": specialties_found,
                "pattern_providers": providers_found,
                "llm_providers": llm_providers,
                "total_providers": len(all_providers),
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Medical provider extraction failed: {str(e)}",
                "file_path": str(file_path),
            }

    async def _extract_key_points(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract key points, arguments, and claims made in a document."""
        file_path = Path(file_path)

        if not file_path.exists():
            return {
                "success": False,
                "error": "File does not exist",
                "file_path": str(file_path),
            }

        try:
            # Extract text content
            content_result = await self._extract_text_content(file_path)
            if not content_result["success"]:
                return {
                    "success": False,
                    "error": f"Could not extract text: {content_result['error']}",
                    "file_path": str(file_path),
                }

            content = content_result["content"]
            key_points = []

            # Pattern-based key point extraction
            key_point_patterns = [
                r"the.*claimant.*alleges.*that.*([^.]+)",
                r"it.*is.*submitted.*that.*([^.]+)",
                r"the.*evidence.*shows.*that.*([^.]+)",
                r"we.*contend.*that.*([^.]+)",
                r"the.*medical.*evidence.*indicates.*([^.]+)",
                r"our.*position.*is.*that.*([^.]+)",
                r"the.*board.*finds.*that.*([^.]+)",
                r"we.*respectfully.*submit.*that.*([^.]+)",
                r"the.*claimant.*argues.*that.*([^.]+)",
                r"in.*our.*opinion.*([^.]+)",
            ]

            for pattern in key_point_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    point = match.strip()
                    if len(point) > 10 and point not in key_points:
                        key_points.append(point)

            # Extract injury-related points
            injury_patterns = [
                r"suffered.*([^.]+.*injury[^.]+)",
                r"injured.*([^.]+.*at.*work[^.]+)",
                r"diagnosis.*of.*([^.]+)",
                r"treated.*for.*([^.]+)",
                r"ongoing.*symptoms.*include.*([^.]+)",
            ]

            injury_points = []
            for pattern in injury_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    point = match.strip()
                    if len(point) > 5:
                        injury_points.append(point)

            # Extract dates and timeline points
            date_patterns = [
                r"on.*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}).*([^.]+)",
                r"since.*(\d{4}).*([^.]+)",
                r"from.*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}).*to.*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}).*([^.]+)",
            ]

            timeline_points = []
            for pattern in date_patterns:
                matches = re.findall(pattern, content, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, tuple) and len(match) >= 2:
                        timeline_points.append(
                            {"date": match[0], "event": match[-1].strip()}
                        )

            # Use LLM for sophisticated key point extraction
            llm_key_points = []
            try:
                llm = await self._get_llm_instance()
                if llm:
                    prompt = f"""
                    Analyze this WCB document and extract the key points, arguments, and claims being made.
                    Focus on:
                    - Main arguments presented
                    - Key medical findings
                    - Important dates and events
                    - Claims about injury causation
                    - Dispute points between parties
                    - Supporting evidence mentioned

                    Document content: {content[:4000]}...

                    Return as JSON:
                    {{
                        "main_arguments": ["argument1", "argument2"],
                        "medical_findings": ["finding1", "finding2"],
                        "key_dates": ["date and event"],
                        "disputed_points": ["dispute1", "dispute2"]
                    }}
                    """

                    response = await llm.ainvoke(prompt)
                    try:
                        import json

                        llm_result = json.loads(response.content)
                        if isinstance(llm_result, dict):
                            llm_key_points = llm_result
                    except:
                        llm_key_points = {}
            except Exception as e:
                logger.warning(f"LLM key point extraction failed: {e}")

            return {
                "success": True,
                "file_path": str(file_path),
                "key_points": key_points,
                "injury_points": injury_points,
                "timeline_points": timeline_points,
                "llm_analysis": llm_key_points,
                "total_points": len(key_points)
                + len(injury_points)
                + len(timeline_points),
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Key point extraction failed: {str(e)}",
                "file_path": str(file_path),
            }

    async def _process_document_with_tracking(
        self, file_path: Union[str, Path], case_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process document with comprehensive WCB tracking."""
        file_path = Path(file_path)

        # First do regular document processing
        result = await self._process_document(file_path)

        if not result["success"] or not self.enable_wcb_tracking:
            return result

        # If case tracking is enabled, add WCB-specific analysis
        try:
            # Identify document source
            source_result = await self._identify_document_source(file_path)

            # Extract medical providers
            providers_result = await self._extract_medical_providers(file_path)

            # Extract key points
            points_result = await self._extract_key_points(file_path)

            # Add to case tracker if one is active
            if self.current_case_tracker:
                document_info = {
                    "file_name": file_path.name,
                    "file_path": str(file_path),
                    "source": source_result.get("source", DocumentSource.UNKNOWN),
                    "submission_date": result["metadata"].get("created_time"),
                    "submission_type": self._determine_submission_type(
                        result["content"]
                    ),
                    "providers": providers_result.get("providers", []),
                    "key_points": points_result.get("key_points", []),
                }

                self.current_case_tracker.add_submission(document_info)

                # Add medical providers to appropriate categories
                for provider in providers_result.get("providers", []):
                    self.current_case_tracker.add_medical_provider(
                        provider, source_result.get("source", DocumentSource.UNKNOWN)
                    )

                # Add key points
                for point in points_result.get("key_points", []):
                    self.current_case_tracker.add_key_point(
                        point,
                        file_path.name,
                        source_result.get("source", DocumentSource.UNKNOWN),
                    )

            # Enhance result with WCB tracking information
            result["wcb_analysis"] = {
                "document_source": source_result,
                "medical_providers": providers_result,
                "key_points": points_result,
                "case_tracking": case_id is not None,
            }

        except Exception as e:
            logger.warning(f"WCB tracking failed for {file_path}: {e}")
            # Don't fail the whole operation if tracking fails
            result["wcb_analysis_error"] = str(e)

        return result

    def _determine_submission_type(self, content: str) -> str:
        """Determine the type of submission based on content."""
        content_lower = content.lower()

        if re.search(r"initial.*claim|first.*report|claim.*application", content_lower):
            return SubmissionType.INITIAL_CLAIM
        elif re.search(
            r"medical.*report|physician.*report|doctor.*report", content_lower
        ):
            return SubmissionType.MEDICAL_REPORT
        elif re.search(
            r"denial.*response|appeal.*submission|challenging.*decision", content_lower
        ):
            return SubmissionType.DENIAL_RESPONSE
        elif re.search(r"appeal|reconsideration|review.*request", content_lower):
            return SubmissionType.APPEAL_SUBMISSION
        elif re.search(r"supporting.*evidence|additional.*evidence", content_lower):
            return SubmissionType.SUPPORTING_EVIDENCE
        elif re.search(r"counter.*argument|rebuttal|response.*to", content_lower):
            return SubmissionType.COUNTER_ARGUMENT
        elif re.search(r"expert.*opinion|independent.*medical", content_lower):
            return SubmissionType.EXPERT_OPINION
        elif re.search(r"correspondence|letter|communication", content_lower):
            return SubmissionType.CORRESPONDENCE
        else:
            return SubmissionType.UNKNOWN
